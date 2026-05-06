"""Génère documentation_faceid.docx dans le répertoire courant."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Styles globaux ────────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(10.5)

for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Calibri'
    h.font.bold = True
    h.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

doc.styles['Heading 1'].font.size = Pt(16)
doc.styles['Heading 2'].font.size = Pt(13)
doc.styles['Heading 3'].font.size = Pt(11)

# Marges
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)


# ── Helpers ───────────────────────────────────────────────────────────────────

def h1(text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)

def h3(text):
    p = doc.add_heading(text, level=3)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)

def para(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent   = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after   = Pt(2)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x16, 0x50, 0x2E)
    # fond gris clair via shading XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F0F4F8')
    pPr.append(shd)
    return p

def table_with_header(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header row
    hrow = t.rows[0]
    for i, hdr in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = hdr
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  '1E3A5F')
        tcPr.append(shd)
    # data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        for ci, val in enumerate(row_data):
            row.cells[ci].text = str(val)
            if ri % 2 == 1:
                tc = row.cells[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'),   'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'),  'EBF0F7')
                tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = w
    doc.add_paragraph()
    return t

def divider():
    doc.add_paragraph('─' * 80).paragraph_format.space_after = Pt(2)

# ═════════════════════════════════════════════════════════════════════════════
#  PAGE DE TITRE
# ═════════════════════════════════════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run('\n\n\nFACEID DASHBOARD')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run('Documentation Fonctionnelle & Technique')
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x55, 0x6B, 0x8D)

p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p_date.add_run(f'\nVersion 1.0  —  {datetime.date.today().strftime("%d %B %Y")}')
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
#  SOMMAIRE
# ═════════════════════════════════════════════════════════════════════════════
h1('Table des matières')
toc_items = [
    ('1', 'Vue d\'ensemble fonctionnelle'),
    ('2', 'Architecture technique'),
    ('3', 'Structure des fichiers'),
    ('4', 'Modules Python — détail'),
    ('  4.1', 'app.py — Serveur Flask & API REST'),
    ('  4.2', 'face_recognizer.py — Reconnaissance faciale'),
    ('  4.3', 'analyzer.py — Analyse démographique'),
    ('  4.4', 'behavior.py — Analyse comportementale'),
    ('  4.5', 'tracker.py — Journal des apparitions'),
    ('  4.6', 'camera_manager.py — Gestion multi-caméras'),
    ('5', 'API REST — Référence complète'),
    ('6', 'Base de données SQLite'),
    ('7', 'Frontend — index.html'),
    ('8', 'Configuration & paramètres'),
    ('9', 'Performance & threads'),
    ('10', 'Dépendances & installation'),
]
for num, title in toc_items:
    p = doc.add_paragraph(f'{num}   {title}')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.2 if num.startswith(' ') else 0)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
#  1. VUE D'ENSEMBLE FONCTIONNELLE
# ═════════════════════════════════════════════════════════════════════════════
h1('1. Vue d\'ensemble fonctionnelle')
para(
    'FaceID Dashboard est une application de vidéosurveillance intelligente accessible par navigateur web. '
    'Elle exploite des modèles de deep learning pour reconnaître des visages, estimer les caractéristiques '
    'démographiques des inconnus, analyser les postures corporelles et détecter les chutes en temps réel.',
)
para(
    'L\'application tourne entièrement en local : aucune donnée n\'est envoyée vers le cloud. '
    'Elle supporte simultanément plusieurs caméras (USB, IP/RTSP, flux HTTP).'
)

h2('1.1 Fonctionnalités principales')
features = [
    ('Tableau de bord (Dashboard)', 'Flux vidéo en direct avec superposition des détections faciales et de posture ; statistiques en temps réel.'),
    ('Analyse faciale (Analyse)', 'Fenêtre caméra plein centre avec bounding boxes colorées, nom, âge estimé et genre.'),
    ('Historique des visages', 'Journal paginé de toutes les apparitions avec filtres caméra, zone, date, statut et recherche par nom.'),
    ('Gestion des personnes', 'Ajout / suppression de profils via upload photo ou capture webcam ; galerie d\'images par personne.'),
    ('Analyse comportementale', 'Détection de posture (debout, assis, allongé, course, saut, grimpe) via MediaPipe Pose.'),
    ('Historique comportemental', 'Chronologie des sessions de posture avec filtre caméra et durée.'),
    ('Détection de chutes', 'Alerte automatique si transition debout/course → allongé en < 2,5 s.'),
    ('Paramètres comportementaux', 'Choix des postures suivies, seuil de confiance, affichage landmarks, caméra surveillée.'),
    ('Gestion des caméras', 'CRUD complet : ajout RTSP/HTTP/webcam, activation/désactivation, association zones.'),
    ('Gestion des utilisateurs', 'Comptes locaux avec rôles admin/viewer ; authentification par session Flask.'),
    ('Paramètres système', 'Seuil de reconnaissance, réinitialisation historique, configuration globale.'),
]
table_with_header(
    ['Fonctionnalité', 'Description'],
    features,
    [Inches(2.2), Inches(4.3)]
)

h2('1.2 Flux utilisateur type')
steps = [
    'L\'utilisateur ouvre http://localhost:5000 et se connecte.',
    'Le serveur démarre tous les CameraWorker actifs (chargés depuis cameras.json).',
    'Chaque worker capte le flux via un SharedCap partagé et effectue la reconnaissance toutes les 4 trames.',
    'Les résultats sont transmis en temps réel via /video_feed (MJPEG) et /api/results (JSON polling).',
    'Les apparitions sont journalisées dans tracker.py (SQLite) et consultables dans l\'onglet Historique.',
    'L\'administrateur gère les profils (personnes/), les caméras et les comptes depuis le menu latéral.',
]
for i, s in enumerate(steps, 1):
    bullet(f'{i}. {s}')

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
#  2. ARCHITECTURE TECHNIQUE
# ═════════════════════════════════════════════════════════════════════════════
h1('2. Architecture technique')
h2('2.1 Vue d\'ensemble')
para('L\'application suit une architecture multi-thread Producer/Consumer :')
arch_lines = [
    'Browser  ←──MJPEG/SSE──→  Flask (app.py)  ←──→  CameraManager',
    '                                                        │',
    '                                    ┌──────────────────┤',
    '                                    │                  │',
    '                             CameraWorker[N]    SharedCap[URL]',
    '                                    │',
    '                    ┌───────────────┼───────────────────┐',
    '                    │               │                   │',
    '             FaceRecognizer    analyzer.py         behavior.py',
    '             (YuNet+SFace)    (DeepFace queue)  (MediaPipe queue)',
    '                    │',
    '               tracker.py (SQLite WAL)',
]
for line in arch_lines:
    code_block(line)

h2('2.2 Technologies utilisées')
table_with_header(
    ['Composant', 'Technologie', 'Rôle'],
    [
        ('Serveur web',       'Flask 2.x',                 'HTTP, MJPEG streaming, API REST, sessions'),
        ('Détection visages', 'OpenCV YuNet (ONNX)',        'Détection rapide en temps réel'),
        ('Reconnaissance',    'OpenCV SFace (ONNX)',        'Embeddings 128D, distance cosinus'),
        ('Démographie',       'DeepFace',                   'Âge et genre sur visages inconnus'),
        ('Posture',           'MediaPipe Pose',             'Squelette 33 landmarks, classification heuristique'),
        ('Persistance',       'SQLite WAL',                 'Journal des apparitions (tracker_log)'),
        ('Config caméras',    'JSON (cameras.json)',        'Persistance CRUD caméras'),
        ('Config users',      'JSON (users.json)',          'Comptes locaux hashés'),
        ('Frontend',          'Tailwind CSS + JS vanilla',  'SPA sans framework'),
        ('Capture vidéo',     'OpenCV VideoCapture',        'SharedCap pool, un thread par URL unique'),
    ],
    [Inches(1.7), Inches(2), Inches(2.8)]
)

h2('2.3 Modèles ONNX')
table_with_header(
    ['Fichier', 'Rôle', 'Localisation'],
    [
        ('face_detection_yunet_2023mar.onnx', 'Détection de visages', 'Racine projet'),
        ('face_recognition_sface_2021dec.onnx', 'Embeddings faciaux',  'Racine projet'),
    ],
    [Inches(3.2), Inches(1.8), Inches(1.5)]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
#  3. STRUCTURE DES FICHIERS
# ═════════════════════════════════════════════════════════════════════════════
h1('3. Structure des fichiers')
tree = """projet/
├── app.py                          # Serveur Flask, routes API, streaming MJPEG
├── camera_manager.py               # CRUD caméras, CameraWorker, SharedCap pool
├── face_recognizer.py              # YuNet + SFace : détection & reconnaissance
├── analyzer.py                     # DeepFace : âge/genre (thread queue)
├── behavior.py                     # MediaPipe Pose : posture & chutes (thread queue)
├── tracker.py                      # Journal SQLite WAL des apparitions
│
├── cameras.json                    # Configuration persistante des caméras
├── users.json                      # Comptes utilisateurs (hashés SHA-256)
├── history.db                      # Base SQLite (table tracker_log)
│
├── known_faces/                    # Profils enregistrés
│   └── <Prénom Nom>/               # Un sous-dossier par personne
│       ├── photo1.jpg
│       └── photo2.jpg
│
├── face_detection_yunet_2023mar.onnx
├── face_recognition_sface_2021dec.onnx
│
└── templates/
    └── index.html                  # SPA complète (Tailwind CSS, JS vanilla)"""
code_block(tree)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
#  4. MODULES PYTHON
# ═════════════════════════════════════════════════════════════════════════════
h1('4. Modules Python — détail')

# ── 4.1 app.py ──────────────────────────────────────────────────────────────
h2('4.1 app.py — Serveur Flask & API REST')
para('Fichier principal. Initialise Flask, démarre tous les sous-systèmes, expose toutes les routes HTTP.')

h3('Initialisation (lignes ~1-80)')
bullet('Création de l\'application Flask')
bullet('Chargement de users.json → dictionnaire en mémoire')
bullet('Instanciation de CameraManager (→ charge cameras.json et démarre les workers)')
bullet('Démarrage de analyzer.start() et behavior.start()')
bullet('Enregistrement du callback analyzer → tracker.update_demographics()')

h3('Authentification')
para('Session Flask avec clé secrète aléatoire. Décorateur @login_required protège toutes les routes sensibles.')
code_block('@login_required\ndef route_protegee():\n    ...')

h3('Routes principales')
table_with_header(
    ['Méthode', 'Route', 'Auth', 'Description'],
    [
        ('GET',        '/',                         'Oui', 'SPA principale (index.html)'),
        ('GET/POST',   '/login',                    'Non', 'Authentification'),
        ('GET',        '/logout',                   'Oui', 'Déconnexion'),
        ('GET',        '/video_feed',               'Oui', 'MJPEG stream caméra active'),
        ('GET',        '/video_feed/<cam_id>',      'Oui', 'MJPEG stream caméra spécifique'),
        ('GET',        '/api/results',              'Oui', 'JSON : résultats reconnaissance en cours'),
        ('GET',        '/api/tracker/log',          'Oui', 'JSON : historique des apparitions'),
        ('GET',        '/api/tracker/stats',        'Oui', 'JSON : total/présents/identifiés/inconnus'),
        ('DELETE',     '/api/tracker/log',          'Oui', 'Suppression sélective ou totale'),
        ('GET',        '/api/cameras',              'Oui', 'Liste des caméras configurées'),
        ('POST',       '/api/cameras',              'Oui', 'Ajout d\'une caméra'),
        ('PUT',        '/api/cameras/<id>',         'Oui', 'Mise à jour caméra'),
        ('DELETE',     '/api/cameras/<id>',         'Oui', 'Suppression caméra'),
        ('POST',       '/api/cameras/<id>/toggle',  'Oui', 'Activer/désactiver'),
        ('GET',        '/api/faces',                'Oui', 'Liste des profils enregistrés'),
        ('GET',        '/face_photos/<name>',       'Oui', 'Photos d\'un profil (JSON base64)'),
        ('POST',       '/add_face',                 'Oui', 'Upload photo de profil'),
        ('POST',       '/capture_face',             'Oui', 'Capture webcam → profil'),
        ('DELETE',     '/delete_face/<name>',       'Oui', 'Suppression d\'un profil'),
        ('GET/POST',   '/api/settings',             'Oui', 'Lecture/écriture seuil de reconnaissance'),
        ('GET',        '/api/behavior/results',     'Oui', 'Posture courante (JSON)'),
        ('GET',        '/api/behavior/history',     'Oui', 'Historique comportemental'),
        ('DELETE',     '/api/behavior/history',     'Oui', 'Effacement historique comportemental'),
        ('GET',        '/api/behavior/falls',       'Oui', 'Historique des chutes'),
        ('DELETE',     '/api/behavior/falls',       'Oui', 'Effacement historique chutes'),
        ('GET/POST',   '/api/behavior/settings',    'Oui', 'Paramètres comportementaux'),
        ('GET',        '/api/users',                'Oui', 'Liste des utilisateurs'),
        ('POST',       '/api/users',                'Oui', 'Création utilisateur'),
        ('DELETE',     '/api/users/<name>',         'Oui', 'Suppression utilisateur'),
        ('POST',       '/api/users/<name>/role',    'Oui', 'Changement de rôle'),
        ('GET',        '/api/rooms',                'Oui', 'Liste des zones disponibles'),
    ],
    [Inches(1.0), Inches(2.5), Inches(0.6), Inches(2.4)]
)

# ── 4.2 face_recognizer.py ──────────────────────────────────────────────────
h2('4.2 face_recognizer.py — Reconnaissance faciale')
para('Encapsule le pipeline YuNet (détection) + SFace (reconnaissance).')

h3('Classe FaceRecognizer')
table_with_header(
    ['Méthode / Attribut', 'Description'],
    [
        ('__init__()',             'Charge les modèles ONNX YuNet et SFace via cv2.FaceDetectorYN et cv2.FaceRecognizerSF'),
        ('threshold',             'Seuil de distance cosinus (défaut 0.363). Modifiable à chaud.'),
        ('load_known_faces()',    'Parcourt known_faces/, encode chaque image, stocke embeddings en mémoire'),
        ('detect_faces(frame)',   'Retourne [(x,y,w,h)] — liste des bounding boxes détectées'),
        ('recognize(frame)',      'Retourne [(x,y,w,h,name)] — détection + identification ou "Inconnu"'),
        ('_match(embedding)',     'Compare l\'embedding à tous les profils, retourne le nom si dist < threshold'),
    ],
    [Inches(2.2), Inches(4.3)]
)

h3('Algorithme de reconnaissance')
steps_rec = [
    'YuNet détecte les visages → bounding boxes + keypoints 5 points',
    'SFace extrait un embedding 128D par visage détecté',
    'Distance cosinus calculée contre tous les embeddings connus',
    'Si dist_min < threshold → nom associé, sinon "Inconnu"',
]
for s in steps_rec:
    bullet(s)

h3('Paramètre clé')
code_block('threshold = 0.363  # face_recognizer.py:__init__\n# Modifiable via /api/settings (POST {threshold: 0.40})')

# ── 4.3 analyzer.py ─────────────────────────────────────────────────────────
h2('4.3 analyzer.py — Analyse démographique')
para(
    'Module optionnel (désactivé si DeepFace absent). Tourne dans un thread dédié avec une queue de taille 1 '
    '(drop silencieux si déjà occupé).'
)

h3('Fonctions publiques')
table_with_header(
    ['Fonction', 'Description'],
    [
        ('start()',                    'Lance le thread worker (no-op si DeepFace manquant ou déjà démarré)'),
        ('submit(frame, faces)',       'Envoie le frame à analyser — uniquement pour les visages "Inconnu"'),
        ('submit_all(frame, faces)',   'Analyse TOUS les visages visibles dans un thread one-shot'),
        ('get_results()',              'Retourne la dernière liste d\'analyses'),
        ('get_all_results()',          'Retourne les résultats de submit_all()'),
        ('set_on_result_callback(fn)', 'Enregistre un callback appelé dès qu\'un résultat est disponible'),
        ('is_available()',             'True si DeepFace est importable'),
    ],
    [Inches(2.5), Inches(4)]
)

h3('Structure d\'un résultat')
code_block("""{
  'bbox':        [x, y, w, h],
  'face':        'data:image/jpeg;base64,...',  # miniature
  'age':         34,
  'age_range':   'Adulte',          # Enfant/Adolescent(e)/Jeune adulte/Adulte/Senior
  'gender':      'Homme',           # Homme | Femme
  'gender_conf': 92.4,              # score confiance %
  'face_size':   'Distance moyenne' # Très proche/Proche/Distance moyenne/Éloigné
}""")

# ── 4.4 behavior.py ─────────────────────────────────────────────────────────
h2('4.4 behavior.py — Analyse comportementale')
para(
    'Détecte la posture en temps réel via MediaPipe Pose (33 landmarks). '
    'Maintient un historique de sessions et détecte les chutes.'
)

h3('Classification des postures — _classify(lm)')
table_with_header(
    ['Posture', 'Heuristique (landmarks MediaPipe)', 'Confiance'],
    [
        ('allonge', 'body_span (cheville - épaule) < 0.22',              '88 %'),
        ('saut',    'cheville.y ≤ hanche.y + 0.08',                      '80 %'),
        ('grimpe',  'poignet.y < nez.y - 0.05',                          '78 %'),
        ('assis',   '(genou.y - hanche.y) < 0.10',                       '82 %'),
        ('course',  'différence genou gauche/droit > 0.11',               '72 %'),
        ('debout',  'aucun critère précédent — cas par défaut',           '90 %'),
        ('inconnu', 'visibilité épaules/hanches < 0.4',                   '50 %'),
    ],
    [Inches(1.2), Inches(3.5), Inches(1.0)]
)

h3('Détection de chutes (_check_fall)')
bullet('Fenêtre glissante de 2,5 secondes (FALL_WINDOW)')
bullet('Chute détectée si : pose actuelle = "allonge" ET une pose "debout" ou "course" dans la fenêtre')
bullet('Enregistrement dans _fall_history avec photo miniature, horodatage et nom')

h3('Structure d\'un événement comportemental')
code_block("""{
  'id':         12,
  'name':       'Alice',
  'photo':      'data:image/jpeg;base64,...',
  'behavior':   'debout',
  'label':      'Debout',
  'color':      'green',
  'first_seen': '2026-05-03 14:22:10',
  'last_seen':  '2026-05-03 14:22:45',
  'duration_s': 35,
  'camera':     'Caméra salon',
}""")

h3('Fonctions publiques')
table_with_header(
    ['Fonction', 'Description'],
    [
        ('start()',               'Lance le worker thread'),
        ('submit(frame, faces, cam_label)', 'Envoie frame à analyser (drop si queue pleine)'),
        ('get_results()',         'Posture courante (liste)'),
        ('get_history(limit)',    'Historique des sessions (max 500)'),
        ('get_fall_history()',    'Historique des chutes (max 100)'),
        ('clear_history()',       'Efface l\'historique comportemental'),
        ('clear_fall_history()', 'Efface l\'historique des chutes'),
        ('get_settings()',        'Retourne show_landmarks, tracked_poses, fall_detect'),
        ('apply_settings(data)',  'Met à jour les paramètres à chaud'),
        ('draw_landmarks_on(frame)', 'Superpose le squelette sur le frame si activé'),
        ('is_available()',        'True si MediaPipe importable'),
    ],
    [Inches(2.8), Inches(3.7)]
)

# ── 4.5 tracker.py ──────────────────────────────────────────────────────────
h2('4.5 tracker.py — Journal des apparitions')
para(
    'Journalise chaque apparition de visage avec début, fin, durée, photo et données démographiques. '
    'Persiste en SQLite (mode WAL pour lectures concurrentes).'
)

h3('Logique de présence')
bullet('Une personne est "présente" si elle a été vue il y a moins de 4 secondes (GONE_AFTER = 4.0)')
bullet('À chaque appel update() : mise à jour last_seen + duration_s pour les actifs, création d\'entrée pour les nouveaux')
bullet('Les entrées "présentes" au redémarrage sont marquées "gone" (nettoyage à l\'init)')

h3('Schéma SQLite — table tracker_log')
code_block("""CREATE TABLE tracker_log (
    id         INTEGER PRIMARY KEY AUTOINCREMENT,
    name       TEXT    NOT NULL DEFAULT '',
    photo      TEXT,                           -- base64 JPEG crop
    first_seen TEXT,                           -- 'YYYY-MM-DD HH:MM:SS'
    last_seen  TEXT,
    duration_s INTEGER DEFAULT 0,
    status     TEXT    DEFAULT 'gone',         -- 'present' | 'gone'
    camera     TEXT,                           -- label de la caméra
    age        INTEGER,                        -- DeepFace estimation
    age_range  TEXT,                           -- tranche d'âge
    gender     TEXT,                           -- 'Homme' | 'Femme'
    face_size  TEXT                            -- distance relative
)""")

h3('Fonctions publiques')
table_with_header(
    ['Fonction', 'Description'],
    [
        ('update(results, frame, analyze_results, cam_label)', 'Mise à jour principale — appelée par CameraWorker'),
        ('update_demographics(analyze_results)',               'Callback DeepFace — enrichit les entrées Inconnu'),
        ('get_log(limit=500)',                                 'Retourne les entrées les plus récentes'),
        ('get_stats()',                                        'Dict total/present/identified/unknown'),
        ('delete_entries(ids)',                                'Suppression sélective par liste d\'IDs'),
        ('clear()',                                            'Effacement complet (mémoire + SQLite)'),
    ],
    [Inches(3.2), Inches(3.3)]
)

# ── 4.6 camera_manager.py ───────────────────────────────────────────────────
h2('4.6 camera_manager.py — Gestion multi-caméras')
para('Deux couches : SharedCap (une VideoCapture par URL) et CameraWorker (logique par caméra).')

h3('SharedCap — pool de captures partagées')
bullet('Un seul thread cv2.VideoCapture par URL unique (même si plusieurs workers pointent vers la même source)')
bullet('subscribe() / unsubscribe() comptent les utilisateurs ; arrêt automatique à 0 utilisateur')
bullet('Stocké dans _shared_caps dict protégé par _sc_lock')

h3('CameraWorker — logique par caméra')
table_with_header(
    ['Attribut / Méthode', 'Description'],
    [
        ('cam_id',             'Identifiant unique (ex: "cam_0", "cam_a3f7b2c1")'),
        ('config',             'Dict de configuration (url, label, rooms, features, ...)'),
        ('_url()',             'Résout l\'URL : int (webcam), RTSP avec credentials injectés, ou string brut'),
        ('start() / stop()',   'Démarre / arrête le thread principal _loop()'),
        ('_loop()',            'Attend SharedCap connecté → lit frames → recognition toutes les 4 trames'),
        ('get_frame()',        'Copie thread-safe du dernier frame'),
        ('get_results()',      'Liste thread-safe des derniers résultats de reconnaissance'),
        ('reload_faces()',     'Recharge les profils connus dans FaceRecognizer'),
        ('to_dict(safe=True)', 'Sérialisation JSON (safe=True masque le mot de passe)'),
    ],
    [Inches(2.2), Inches(4.3)]
)

h3('CameraManager — API de haut niveau')
table_with_header(
    ['Méthode', 'Description'],
    [
        ('list()',                    'Liste toutes les caméras (sérialisées)'),
        ('get(cam_id)',               'Détails d\'une caméra'),
        ('add(data)',                 'Crée et démarre un nouveau worker, persiste en JSON'),
        ('update(cam_id, data)',      'Arrête, reconfigure et redémarre le worker'),
        ('delete(cam_id)',            'Arrête et supprime le worker + persiste'),
        ('toggle(cam_id, enabled)',   'Active ou désactive une caméra'),
        ('reload_all_faces()',        'Recharge les profils dans tous les workers'),
        ('set_threshold(value)',      'Propagation du seuil à tous les FaceRecognizer'),
        ('get_frame(cam_id)',         'Dernière frame d\'une caméra'),
        ('get_results(cam_id)',       'Derniers résultats d\'une caméra'),
        ('first_id() / all_ids()',    'IDs des workers actifs'),
    ],
    [Inches(2.5), Inches(4)]
)

h3('Configuration d\'une caméra (cameras.json)')
code_block("""{
  "id":       "cam_0",
  "label":    "Caméra Entrée",
  "rooms":    ["Entrée"],
  "url":      "0",                 // int webcam, IP, RTSP, HTTP
  "ip":       "",
  "port":     "",
  "protocol": "webcam",            // webcam | rtsp | http
  "username": "",
  "password": "",
  "enabled":  true,
  "features": ["face_recognition", "behavior_analysis"]
}""")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
#  5. API REST — RÉFÉRENCE COMPLÈTE
# ═════════════════════════════════════════════════════════════════════════════
h1('5. API REST — Référence complète')
para('Toutes les routes nécessitent une session authentifiée (cookie de session Flask), sauf /login.')

h2('5.1 Caméras')
code_block('GET  /api/cameras            → [{id, label, rooms, url, enabled, connected, running, error, ...}]')
code_block('POST /api/cameras            ← {label, url, rooms, features, ...}   → {id: "cam_xxx"}')
code_block('PUT  /api/cameras/<id>       ← {label, url, rooms, features, ...}   → {ok: true}')
code_block('DEL  /api/cameras/<id>                                               → {ok: true}')
code_block('POST /api/cameras/<id>/toggle ← {enabled: true|false}               → {ok: true}')

h2('5.2 Flux vidéo')
code_block('GET /video_feed              → multipart/x-mixed-replace (MJPEG) — caméra active courante')
code_block('GET /video_feed/<cam_id>     → multipart/x-mixed-replace (MJPEG) — caméra spécifique')

h2('5.3 Reconnaissance & historique')
code_block('GET /api/results             → [{name, bbox:[x,y,w,h], ...}]        (polling ~500 ms)')
code_block('GET /api/tracker/log        → [{id, name, photo, first_seen, last_seen, duration_s, status, camera, age, age_range, gender, face_size}]')
code_block('GET /api/tracker/stats      → {total, present, identified, unknown}')
code_block('DEL /api/tracker/log        ← {ids:[1,2,3]}  ou  {}  (effacement total)')

h2('5.4 Profils (personnes)')
code_block('GET  /api/faces              → [{name, count, preview}]')
code_block('GET  /face_photos/<name>     → [{filename, data: "data:image/jpeg;base64,..."}]')
code_block('POST /add_face               ← multipart: name + file   → {ok: true}')
code_block('POST /capture_face           ← {name, cam_id}           → {ok: true}')
code_block('DEL  /delete_face/<name>                                → {ok: true}')

h2('5.5 Paramètres')
code_block('GET  /api/settings           → {threshold: 0.363}')
code_block('POST /api/settings           ← {threshold: 0.40}        → {ok: true}')

h2('5.6 Comportemental')
code_block('GET  /api/behavior/results   → [{pose, label, color, confidence, name, vid_text}]')
code_block('GET  /api/behavior/history   → [{id, name, photo, behavior, label, color, first_seen, last_seen, duration_s, camera}]')
code_block('DEL  /api/behavior/history                               → {ok: true}')
code_block('GET  /api/behavior/falls     → [{id, name, timestamp, photo}]')
code_block('DEL  /api/behavior/falls                                 → {ok: true}')
code_block('GET  /api/behavior/settings  → {show_landmarks, tracked_poses:[], fall_detect}')
code_block('POST /api/behavior/settings  ← {show_landmarks, tracked_poses:[], fall_detect}  → {ok: true}')

h2('5.7 Utilisateurs')
code_block('GET  /api/users              → [{username, role}]')
code_block('POST /api/users              ← {username, password, role}  → {ok: true}')
code_block('DEL  /api/users/<name>                                     → {ok: true}')
code_block('POST /api/users/<name>/role  ← {role: "admin"|"viewer"}   → {ok: true}')

h2('5.8 Zones')
code_block("GET  /api/rooms              → ['Entrée','Salon','Cuisine',...]")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
#  6. BASE DE DONNÉES
# ═════════════════════════════════════════════════════════════════════════════
h1('6. Base de données SQLite')
para('Fichier : history.db — mode WAL (Write-Ahead Logging) pour lectures concurrentes.')

table_with_header(
    ['Colonne', 'Type', 'Description'],
    [
        ('id',         'INTEGER PK', 'Auto-incrémenté'),
        ('name',       'TEXT',       'Nom reconnu ou "Inconnu"'),
        ('photo',      'TEXT',       'data:image/jpeg;base64,... (crop visage)'),
        ('first_seen', 'TEXT',       'Horodatage première apparition (YYYY-MM-DD HH:MM:SS)'),
        ('last_seen',  'TEXT',       'Horodatage dernière détection'),
        ('duration_s', 'INTEGER',    'Durée totale de présence en secondes'),
        ('status',     'TEXT',       '"present" | "gone" (gone après 4s d\'absence)'),
        ('camera',     'TEXT',       'Label de la caméra qui a capté la personne'),
        ('age',        'INTEGER',    'Âge estimé par DeepFace (NULL si non analysé)'),
        ('age_range',  'TEXT',       'Tranche : Enfant / Adolescent(e) / Jeune adulte / Adulte / Senior'),
        ('gender',     'TEXT',       '"Homme" | "Femme" (NULL si non analysé)'),
        ('face_size',  'TEXT',       'Distance estimée : Très proche / Proche / Distance moyenne / Éloigné'),
    ],
    [Inches(1.2), Inches(1.1), Inches(4.2)]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
#  7. FRONTEND
# ═════════════════════════════════════════════════════════════════════════════
h1('7. Frontend — templates/index.html')
para('SPA (Single Page Application) sans framework JS. Navigation simulée par masquage/affichage de divs.')

h2('7.1 Zones d\'affichage (divs principales)')
table_with_header(
    ['ID', 'Visible quand', 'Contenu'],
    [
        ('main-default',    'Dashboard, Analyse, Comportement, Paramètres', 'Flux caméra MJPEG principal'),
        ('main-add',        'nav = "add"',           'Aperçu caméra centré pour capture de profil'),
        ('main-people',     'nav = "people"',        'Galerie photos d\'un profil sélectionné'),
        ('main-history',    'nav = "history"',       'Tableau historique des apparitions'),
        ('main-behavhist',  'nav = "behavhist"',     'Tableau historique comportemental'),
        ('main-cameras',    'nav = "cameras"',       'Tableau de gestion des caméras'),
        ('main-users',      'nav = "users"',         'Tableau de gestion des utilisateurs'),
    ],
    [Inches(1.7), Inches(2.3), Inches(2.5)]
)

h2('7.2 Panneaux droits (aside)')
table_with_header(
    ['ID', 'Associé à nav', 'Contenu'],
    [
        ('panel-dashboard',    'dashboard',    'Statistiques live, miniatures personnes présentes'),
        ('panel-analyse',      'analyse',      'Sélecteur caméra, seuil, toggle landmarks'),
        ('panel-history',      'history',      'Filtres caméra/zone/date/statut, recherche'),
        ('panel-add',          'add',          'Formulaire "Nouveau profil" (nom, upload, capture)'),
        ('panel-people',       'people',       'Liste des profils enregistrés'),
        ('panel-settings',     'settings',     'Seuil global, reset historique, gestion users inline'),
        ('panel-behavior',     'behavior',     'Résultats posture live, historique session'),
        ('panel-behavhist',    'behavhist',    'Filtre caméra historique comportemental'),
        ('panel-behavsettings','behavsettings','Caméra surveillée, landmarks, chutes, postures'),
        ('panel-cameras',      'cameras',      'Formulaire ajout/édition caméra'),
        ('panel-users',        'users',        'Formulaire ajout utilisateur'),
    ],
    [Inches(2), Inches(1.8), Inches(2.7)]
)

h2('7.3 Fonctions JavaScript clés')
table_with_header(
    ['Fonction', 'Rôle'],
    [
        ('setNav(name)',              'Routing principal : affiche le bon main-* et panel-*'),
        ('toggleRightPanel()',        'Replie/déplie l\'aside droit (transition CSS width)'),
        ('refreshStats()',            'Met à jour les compteurs dashboard via /api/tracker/stats'),
        ('loadFaces()',               'Charge la liste des profils depuis /api/faces'),
        ('openPersonGallery(name)',   'Affiche la galerie photos d\'un profil dans main-people'),
        ('renderBehaviorHistory()',   'Rend le tableau comportemental avec filtres caméra/posture'),
        ('setBehavCamFilter(val)',    'Filtre l\'historique comportemental par caméra'),
        ('_populateCamSelectors()',   'Remplit tous les <select> de caméras'),
        ('onAddCamChange()',          'Charge l\'aperçu caméra dans main-add'),
        ('refreshBehaviorResults()',  'Polling /api/behavior/results → affichage posture live'),
        ('loadHistory()',             'Charge l\'historique des apparitions depuis /api/tracker/log'),
        ('loadCameras()',             'Charge et affiche la liste des caméras'),
        ('submitCameraForm()',        'Soumet le formulaire d\'ajout/édition caméra'),
        ('deleteCamera(id)',          'Supprime une caméra avec confirmation'),
        ('addFace()',                 'Upload photo de profil via /add_face'),
        ('captureFace()',             'Capture webcam → /capture_face'),
        ('deleteFace(name)',          'Supprime un profil avec confirmation'),
    ],
    [Inches(2.5), Inches(4)]
)

h2('7.4 Polling & intervals')
table_with_header(
    ['Donnée', 'Intervalle', 'Route'],
    [
        ('Statistiques dashboard',   '3 000 ms', '/api/tracker/stats'),
        ('Résultats reconnaissance', '500 ms',   '/api/results'),
        ('Résultats comportementaux','1 000 ms', '/api/behavior/results'),
        ('Flux MJPEG',               'continu',  '/video_feed ou /video_feed/<id>'),
    ],
    [Inches(2.2), Inches(1.3), Inches(3)]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
#  8. CONFIGURATION & PARAMÈTRES
# ═════════════════════════════════════════════════════════════════════════════
h1('8. Configuration & paramètres')
table_with_header(
    ['Paramètre', 'Valeur défaut', 'Fichier / ligne', 'Description'],
    [
        ('threshold',    '0.363',    'face_recognizer.py:__init__',  'Seuil distance cosinus SFace'),
        ('GONE_AFTER',   '4.0 s',    'tracker.py:14',                'Délai avant marquage "gone"'),
        ('MAX_LOG',      '500',      'tracker.py:15',                'Entrées max en mémoire'),
        ('FALL_WINDOW',  '2.5 s',    'behavior.py:55',               'Fenêtre détection chute'),
        ('MAX_HISTORY',  '500',      'behavior.py:48',               'Max sessions comportementales'),
        ('CAMERAS_FILE', 'cameras.json', 'camera_manager.py:15',     'Config persistante caméras'),
        ('DB_FILE',      'history.db',   'tracker.py:18',             'Base SQLite'),
        ('maxsize queue','1',        'analyzer.py:17 / behavior.py:31', 'Queue drop-if-full'),
        ('frame skip',   'idx%4',    'camera_manager.py:249',        'Reconnaissance 1 frame/4'),
        ('connect timeout','60 s',   'camera_manager.py:209',        'Timeout connexion SharedCap'),
    ],
    [Inches(1.6), Inches(1.2), Inches(2.2), Inches(1.5)]
)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
#  9. PERFORMANCE & THREADS
# ═════════════════════════════════════════════════════════════════════════════
h1('9. Performance & threads')
h2('9.1 Threads actifs en fonctionnement normal')
table_with_header(
    ['Thread', 'Fichier', 'Rôle'],
    [
        ('SharedCap._capture_loop × N_URL', 'camera_manager.py', 'Capture vidéo (un par URL unique)'),
        ('CameraWorker._loop × N_CAM',      'camera_manager.py', 'Reconnaissance + orchestration'),
        ('analyzer._worker',                'analyzer.py',       'Analyse DeepFace (queue 1)'),
        ('behavior._worker',                'behavior.py',       'Analyse MediaPipe (queue 1)'),
        ('Flask request threads',           'app.py',            'Requêtes HTTP concurrentes'),
    ],
    [Inches(2.5), Inches(1.8), Inches(2.2)]
)

h2('9.2 Optimisations')
bullet('SharedCap évite les doublons VideoCapture si deux workers pointent vers la même caméra')
bullet('Queue(maxsize=1) + put_nowait() : drop silencieux si le worker précédent n\'a pas fini (pas de backpressure)')
bullet('frame skip × 4 : reconnaissance toutes les ~133 ms à 30 fps → charge CPU réduite')
bullet('SQLite WAL : lectures non bloquantes pendant les écritures')
bullet('Embeddings chargés une seule fois en mémoire au démarrage (load_known_faces)')

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
#  10. DÉPENDANCES & INSTALLATION
# ═════════════════════════════════════════════════════════════════════════════
h1('10. Dépendances & installation')
h2('10.1 Dépendances Python')
table_with_header(
    ['Package', 'Usage'],
    [
        ('flask',           'Serveur web, sessions, routing'),
        ('opencv-python',   'Capture vidéo, YuNet, SFace, encodage JPEG'),
        ('mediapipe',       'Pose estimation (optionnel)'),
        ('deepface',        'Analyse démographique (optionnel)'),
        ('numpy',           'Calculs vectoriels'),
        ('tensorflow / tf-intel', 'Backend DeepFace (optionnel)'),
        ('python-docx',     'Génération de ce document (utilitaire ponctuel)'),
    ],
    [Inches(2.2), Inches(4.3)]
)

h2('10.2 Lancement')
code_block('# Installer les dépendances\npip install flask opencv-python mediapipe deepface\n\n# Lancer le serveur\npython app.py\n\n# Accéder à l\'interface\nhttp://localhost:5000')

h2('10.3 Identifiants par défaut')
code_block('Utilisateur : admin\nMot de passe : admin\n\n# Modifiable via Interface → Paramètres → Gestion des utilisateurs')

h2('10.4 Ajout d\'un profil facial')
steps_face = [
    'Créer un dossier known_faces/<Prénom Nom>/',
    'Placer au moins 3 photos JPEG/PNG claires (visage bien visible)',
    'Dans l\'interface : Paramètres → Recharger les visages  OU  python app.py (redémarrage)',
    'Alternative : utiliser l\'onglet "Ajouter personne" → upload ou capture webcam',
]
for s in steps_face:
    bullet(s)

# ── Pied de page ─────────────────────────────────────────────────────────────
doc.add_page_break()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = p_footer.add_run(
    f'FaceID Dashboard — Documentation v1.0\n'
    f'Générée le {datetime.date.today().strftime("%d %B %Y")}\n'
    f'Repo : https://github.com/marouane1980/reconnaissance-faciale'
)
rf.font.size = Pt(9)
rf.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Sauvegarde ────────────────────────────────────────────────────────────────
out_path = r'C:\Users\us\project\documentation_faceid.docx'
doc.save(out_path)
print(f'[OK] Document sauvegardé : {out_path}')
